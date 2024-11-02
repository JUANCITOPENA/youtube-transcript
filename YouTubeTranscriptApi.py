# Para instalar las herramientas necesarias: 
# pip install streamlit youtube-transcript-api python-docx

# Importar las bibliotecas necesarias
import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi  # Para obtener transcripciones de videos de YouTube
from docx import Document  # Para crear y manipular documentos de Word
import io  # Para manejo de flujos de entrada y salida en memoria

# Título de la aplicación, usando HTML para estilo de encabezado
st.markdown("<h3 style='text-align: center;'>🎥 Transcribe el contenido de tu video a Texto en Word ⭐</h3>", unsafe_allow_html=True)

# Banner debajo del título, con una imagen centralizada desde un URL de repositorio
st.markdown(
    "<div style='text-align: center;'><img src='https://raw.githubusercontent.com/JUANCITOPENA/-REMOVER-FONDO-IM-GENES-/refs/heads/main/banner.webp' alt='Banner' style='width: 100%;'></div>", 
    unsafe_allow_html=True
)

# Campo de texto para que el usuario ingrese la URL del video de YouTube
video_url = st.text_input("Ingresa la URL del video de YouTube:")

# Verificar si el botón "Generar Transcripción" fue presionado
if st.button("Generar Transcripción"):
    # Verificar si la URL del video fue ingresada
    if not video_url:
        st.error("Por favor, ingresa una URL de video.")  # Mostrar mensaje de error si el campo está vacío
    else:
        try:
            # Extraer el ID del video de YouTube de la URL ingresada
            video_id = video_url.split("v=")[-1]
            # Obtener la transcripción del video en español usando la API de YouTubeTranscriptApi
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['es'])

            # Crear un nuevo documento de Word
            doc = Document()
            doc.add_heading('Transcripción del Video', level=1)  # Añadir un encabezado al documento

            # Agregar cada parte de la transcripción al documento de Word
            for entry in transcript:
                doc.add_paragraph(entry['text'])  # Insertar el texto de cada entrada de la transcripción como un párrafo

            # Guardar el documento en un objeto BytesIO para ser descargado sin almacenarlo en disco
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)  # Establecer el puntero al inicio del archivo en memoria

            # Crear un botón de descarga para el documento de Word generado
            st.download_button(
                label="Descargar Transcripción",
                data=output,
                file_name="transcripcion_video.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.success("La transcripción está lista para descargarse.")  # Mensaje de éxito al generar la transcripción

        except Exception as e:
            # Capturar y mostrar un error si la transcripción no pudo generarse
            st.error(f"No se pudo obtener la transcripción: {e}")
